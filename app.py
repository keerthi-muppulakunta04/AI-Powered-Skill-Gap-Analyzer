# app.py
# 🌟 AI Powered Skill Gap Analyzer for Resume and Job Description Matching 🌟
# Student Name: Keerthi Muppulakunta
# Student ID: 17
import os
import re
import io
import json
import tempfile
import logging
from io import BytesIO
from typing import List, Dict, Any, Tuple, Optional
from datetime import datetime

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go

# NLP & Embeddings
import spacy
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity

# File parsing & OCR
import PyPDF2
import docx
try:
    from pdf2image import convert_from_bytes
    pdf2image_available = True
except Exception:
    convert_from_bytes = None
    pdf2image_available = False

try:
    import pytesseract
    pytesseract_available = True
except Exception:
    pytesseract = None
    pytesseract_available = False

# Fuzzy matching
from fuzzywuzzy import fuzz

# PDF report
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    reportlab_available = True
except Exception:
    reportlab_available = False

# Optional: save plotly images (kaleido)
try:
    import kaleido  # type: ignore # noqa: F401
    kaleido_available = True
except Exception:
    kaleido_available = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("SkillGapApp")
 
# ---------------- UI config ----------------
st.set_page_config(page_title="AI Skill Gap Analyzer", layout="wide", page_icon="🌟")

# Light blue theme + banner
st.markdown(
    """
    <style>
    [data-testid="stSidebar"] > div:first-child {
        background: linear-gradient(180deg,#e3f6ff 0%,#f0fbff 100%);
        color: #002b5c;
    }
    .stApp {
        background: linear-gradient(180deg,#f6fdff 0%,#ffffff 100%);
    }
    .banner {
        background: linear-gradient(90deg,#d8f2ff,#eaf9ff);
        color: #002b5c;
        padding: 14px;
        border-radius: 10px;
        text-align: center;
        font-size: 20px;
        font-weight: 700;
        margin-bottom: 12px;
    }
    .chip {
        display:inline-block;
        padding:6px 10px;
        border-radius:12px;
        background:#e6f7ff;
        color:#003147;
        font-weight:600;
        margin:4px;
    }
    .skill-btn {
        display:inline-block;
        margin:4px;
        padding:8px 12px;
        border-radius:12px;
        font-weight:600;
        cursor: pointer;
    }
    .skill-strong { background:#79d279; color:#003200; }
    .skill-partial { background:#ffe082; color:#664d00; }
    .skill-missing { background:#ff9a9a; color:#660000; }
    </style>
    """, unsafe_allow_html=True
)

st.markdown("<div class='banner'>🌟 AI Powered Skill Gap Analyzer for Resume and Job Description Matching 🌟</div>", unsafe_allow_html=True)

# ---------------- Helper classes ----------------

class DocumentProcessor:
    SUPPORTED = ("pdf", "docx", "txt")
    MAX_SIZE = 100 * 1024 * 1024  # 100MB

    def __init__(self, tesseract_cmd: Optional[str] = None):
        # Optionally configure tesseract path
        if tesseract_cmd and pytesseract:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    def validate_file(self, f) -> Tuple[bool, Optional[str]]:
        if f is None:
            return False, "No file"
        if f.size == 0:
            return False, "Empty file"
        ext = f.name.split('.')[-1].lower()
        if ext not in self.SUPPORTED:
            return False, f"Unsupported: {ext}"
        if f.size > self.MAX_SIZE:
            return False, "File too large"
        return True, None

    def extract_text(self, uploaded_file) -> Dict[str, Any]:
        name = uploaded_file.name
        ext = name.split('.')[-1].lower()
        raw_text = ""
        pages = None
        parse_status = "failed"
        error = None
        ocr_conf = None

        try:
            if ext == 'pdf':
                try:
                    reader = PyPDF2.PdfReader(BytesIO(uploaded_file.getvalue()))
                    pages = len(reader.pages)
                    for i, page in enumerate(reader.pages):
                        try:
                            txt = page.extract_text() or ""
                            raw_text += f"\n--- Page {i+1} ---\n" + txt
                        except Exception:
                            continue
                    parse_status = "success" if raw_text.strip() else "empty"
                except Exception as e:
                    # fallback to raw bytes read attempt
                    error = f"PyPDF2 read error: {e}"
                    parse_status = "failed"

                # OCR fallback if empty and dependencies present
                if (not raw_text.strip()) and pdf2image_available and pytesseract_available:
                    try:
                        imgs = convert_from_bytes(uploaded_file.getvalue(), dpi=200)
                        ocr_text = ""
                        confidences = []
                        for img in imgs:
                            data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
                            ocr_text += "\n" + pytesseract.image_to_string(img)
                            # compute average confidence for non-empty words
                            confs = [int(x) for x in data['conf'] if x.isdigit() or (isinstance(x, (int,float)))]
                            if confs:
                                confidences.extend(confs)
                        raw_text = ocr_text
                        parse_status = "ocr" if raw_text.strip() else "failed"
                        if confidences:
                            ocr_conf = float(np.mean(confidences)) / 100.0
                    except Exception as e:
                        logger.exception("OCR fallback failed")
                        error = f"OCR failed: {e}"
            elif ext == 'docx':
                try:
                    doc = docx.Document(BytesIO(uploaded_file.getvalue()))
                    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                    # tables too
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                            if row_text:
                                paragraphs.append(row_text)
                    raw_text = "\n".join(paragraphs)
                    parse_status = "success" if raw_text.strip() else "empty"
                except Exception as e:
                    error = f"DOCX read error: {e}"
            elif ext == 'txt':
                try:
                    raw_text = uploaded_file.getvalue().decode('utf-8')
                except Exception:
                    raw_text = uploaded_file.getvalue().decode('latin-1', errors='replace')
                parse_status = "success"
            else:
                error = f"Unsupported ext: {ext}"
                parse_status = "failed"

        except Exception as e:
            logger.exception("File extraction exception")
            error = str(e)
            parse_status = "failed"

        # Clean & normalize
        cleaned_text, norm_summary = TextCleaner().clean_text(raw_text)
        tokens = len(cleaned_text.split())

        return {
            "name": name,
            "format": ext,
            "raw_text": raw_text,
            "clean_text": cleaned_text,
            "pages": pages,
            "parse_status": parse_status,
            "error": error,
            "tokens": tokens,
            "normalization_summary": norm_summary,
            "ocr_confidence": ocr_conf
        }

class TextCleaner:
    def __init__(self):
        pass

    def clean_text(self, raw_text: str) -> Tuple[str, Dict[str, Any]]:
        if not raw_text:
            return "", {"original_chars":0,"final_chars":0,"removed_lines":0,"redacted_contact_info":False}
        original_chars = len(raw_text)
        text = raw_text.replace('\r\n','\n').replace('\r','\n')
        # remove multiple blank lines
        text = re.sub(r'\n{3,}', '\n\n', text)
        # remove common headers/footers
        text = re.sub(r'Page\s*\d+\s*of\s*\d+', '', text, flags=re.IGNORECASE)
        # remove mocked lines like --- Page x ---
        text = re.sub(r'^---+\s*Page.*$', '', text, flags=re.MULTILINE)
        # redact emails and phones
        redacted = False
        text, n_email = re.subn(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b', '[REDACTED_EMAIL]', text)
        text, n_phone = re.subn(r'(\+?\d[\d\-\s]{7,}\d)', '[REDACTED_PHONE]', text)
        if n_email + n_phone > 0:
            redacted = True
        # standardize bullets/quotes
        text = text.replace('•','-').replace('“','"').replace('”','"').strip()
        final_chars = len(text)
        removed_lines = max(0, (original_chars - final_chars)//10)
        summary = {"original_chars":original_chars, "final_chars":final_chars, "removed_lines":removed_lines, "redacted_contact_info":redacted}
        return text, summary

# ---------------- Skill extraction classes ----------------

class SkillDatabase:
    def __init__(self):
        self.db = {
            "programming": ["Python","Java","C++","C#","JavaScript","R","Go","Swift"],
            "data": ["SQL","PostgreSQL","MySQL","MongoDB","Pandas","NumPy","Scikit-learn","TensorFlow","PyTorch","Spark"],
            "web": ["React","Angular","Django","Flask","Node.js","Express","Vue"],
            "cloud": ["AWS","Azure","GCP","Docker","Kubernetes"],
            "soft": ["Communication","Leadership","Teamwork","Problem Solving","Time Management"]
        }

    def all_skills(self) -> List[str]:
        skills = []
        for lst in self.db.values():
            skills.extend(lst)
        return skills

    def category(self, skill: str) -> str:
        s = skill.strip().lower()
        for k,v in self.db.items():
            for item in v:
                if item.strip().lower() == s:
                    return k
        return "other"

class SkillExtractor:
    def __init__(self, use_spacy=True, use_keyword=True, use_fuzzy=True):
        self.use_spacy = use_spacy
        self.use_keyword = use_keyword
        self.use_fuzzy = use_fuzzy
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except Exception:
            spacy.cli.download("en_core_web_sm")
            self.nlp = spacy.load("en_core_web_sm")
        self.db = SkillDatabase()
        self.master = self.db.all_skills()

    def extract(self, text: str) -> List[Dict[str,Any]]:
        if not text or not text.strip():
            return []
        doc = self.nlp(text[:200000])
        sentences = [s.text.strip() for s in doc.sents]
        found = {}
        # keyword matching
        if self.use_keyword:
            for sk in self.master:
                pat = re.compile(r'\b' + re.escape(sk) + r'\b', re.IGNORECASE)
                matches = pat.findall(text)
                if matches:
                    found.setdefault(sk, {"count":0,"sentences":[]})
                    found[sk]["count"] += len(matches)
                    for s in sentences:
                        if pat.search(s):
                            found[sk]["sentences"].append(s)
        # noun-chunks
        if self.use_spacy:
            for chunk in doc.noun_chunks:
                chunk_text = chunk.text.strip()
                if len(chunk_text.split()) <= 4:
                    for sk in self.master:
                        if sk.lower() in chunk_text.lower():
                            found.setdefault(sk, {"count":0,"sentences":[]})
                            found[sk]["count"] += 1
                            found[sk]["sentences"].append(chunk.sent.text.strip())
        # fuzzy on entities
        if self.use_fuzzy:
            for ent in doc.ents:
                ent_text = ent.text.strip()
                for sk in self.master:
                    score = fuzz.partial_ratio(ent_text.lower(), sk.lower())
                    if score >= 85:
                        found.setdefault(sk, {"count":0,"sentences":[]})
                        found[sk]["count"] += 1
                        found[sk]["sentences"].append(ent.sent.text.strip())

        skills = []
        for sk, info in found.items():
            freq = info["count"]
            conf = min(0.95, 0.3 + 0.15 * np.log1p(freq))
            skills.append({
                "skill": sk,
                "category": self.db.category(sk),
                "count": freq,
                "sentences": list(dict.fromkeys(info["sentences"]))[:5],
                "confidence": round(conf,2)
            })
        skills = sorted(skills, key=lambda x: (x["confidence"], x["count"]), reverse=True)
        return skills

# ---------------- Embeddings & similarity ----------------

class SentenceBERTEmbedder:
    def __init__(self, model_name="all-MiniLM-L6-v2"):
        self.model_name = model_name
        self.model = SentenceTransformer(model_name)
        self.dim = self.model.get_sentence_embedding_dimension()
        self.cache = {}

    def encode(self, texts: List[str], use_cache=True) -> np.ndarray:
        if not texts:
            return np.zeros((0,self.dim))
        to_encode = []
        for t in texts:
            if use_cache and t in self.cache:
                continue
            to_encode.append(t)
        if to_encode:
            emb = self.model.encode(to_encode, show_progress_bar=False)
            for t,e in zip(to_encode, emb):
                self.cache[t] = e
        # return array in same order
        return np.array([self.cache.get(t, np.zeros(self.dim)) for t in texts])

    def clear_cache(self):
        self.cache = {}

class SimilarityUtils:
    @staticmethod
    def cosine_matrix(a: np.ndarray, b: np.ndarray) -> np.ndarray:
        if a.size == 0 or b.size == 0:
            return np.zeros((a.shape[0], b.shape[0]))
        return cosine_similarity(a, b)

# ---------------- Analysis & Visualizer ----------------

class SkillGapAnalyzer:
    def __init__(self, embedder: SentenceBERTEmbedder, strong=0.8, partial=0.5):
        self.embedder = embedder
        self.strong = strong
        self.partial = partial

    def _dedupe(self, skills: List[str], thresh=90) -> List[str]:
        unique = []
        for s in skills:
            s_clean = s.strip()
            if not any(fuzz.ratio(s_clean.lower(), u.lower()) >= thresh for u in unique):
                unique.append(s_clean)
        return unique

    def analyze(self, resume_skills: List[str], jd_skills: List[str], importance: Optional[Dict[str,float]] = None) -> Dict[str,Any]:
        resume = self._dedupe(resume_skills)
        jd = self._dedupe(jd_skills)
        emb_r = self.embedder.encode(resume)
        emb_j = self.embedder.encode(jd)
        sim = SimilarityUtils.cosine_matrix(emb_r, emb_j) if emb_r.size and emb_j.size else np.zeros((len(resume), len(jd)))
        matched, partial, missing = [], [], []
        for j_idx, j_skill in enumerate(jd):
            if sim.size == 0:
                best_sim = 0.0
                best_res = ""
            else:
                best_idx = int(np.argmax(sim[:, j_idx]))
                best_sim = float(sim[best_idx, j_idx])
                best_res = resume[best_idx] if resume else ""
            imp = importance.get(j_skill,1.0) if importance else 1.0
            rec = {"jd_skill": j_skill, "resume_skill": best_res, "similarity": best_sim, "importance": imp}
            if best_sim >= self.strong:
                matched.append(rec)
            elif best_sim >= self.partial:
                partial.append(rec)
            else:
                missing.append(rec)
        overall = float(np.mean(sim.max(axis=0))) if sim.size else 0.0
        return {
            "resume_unique": resume,
            "jd_unique": jd,
            "similarity_matrix": sim,
            "matched": matched,
            "partial": partial,
            "missing": missing,
            "overall": overall
        }

class Visualizer:
    @staticmethod
    def pie(result):
        vals = [len(result["matched"]), len(result["partial"]), len(result["missing"])]
        labels = ["Matched","Partial","Missing"]
        fig = px.pie(values=vals, names=labels, hole=0.35, color_discrete_sequence=["#79d279","#ffe082","#ff9a9a"])
        fig.update_traces(textposition='inside', textinfo='percent+label')
        return fig

    @staticmethod
    def bubble_tagcloud(skills_list: List[Dict[str,Any]], title="Top Skills (bubble size = frequency)"):
        # skills_list: list of dicts with 'skill' and 'count'
        if not skills_list:
            return go.Figure()
        df = pd.DataFrame(skills_list)
        df["size"] = df["count"].clip(lower=1)
        df["x"] = np.random.uniform(0,1,len(df))
        df["y"] = np.random.uniform(0,1,len(df))
        fig = px.scatter(df, x="x", y="y", size="size", text="skill", size_max=60, title=title)
        fig.update_traces(textposition='middle center', marker=dict(opacity=0.8))
        fig.update_xaxes(visible=False)
        fig.update_yaxes(visible=False)
        fig.update_layout(height=400)
        return fig

    @staticmethod
    def bar_missing(result, top_n=15):
        missing = sorted(result["missing"], key=lambda x: (x["importance"]*(1-x["similarity"])))[:top_n]
        if not missing:
            return go.Figure()
        df = pd.DataFrame(missing)
        df["sim_pct"] = df["similarity"]*100
        fig = go.Figure(go.Bar(x=df["sim_pct"][::-1], y=df["jd_skill"][::-1], orientation='h', marker_color="#b0e0ff"))
        fig.update_layout(title="Top Missing Skills (lower similarity = higher gap)", xaxis_title="Similarity (%)")
        return fig

    @staticmethod
    def heatmap(sim_matrix, resume_skills, jd_skills, top=40):
        if sim_matrix.size == 0:
            return go.Figure()
        r = resume_skills[:top]
        j = jd_skills[:top]
        mat = sim_matrix[:len(r), :len(j)]
        fig = go.Figure(data=go.Heatmap(z=mat, x=j, y=r, colorscale="RdYlGn", zmin=0, zmax=1))
        fig.update_layout(title=f"Similarity Heatmap (top {len(r)} x {len(j)})", height=600)
        return fig

    @staticmethod
    def radar_category(result, resume_skills, jd_skills, extractor: SkillExtractor):
        # compute category-wise average similarity
        cats = list(set(list(extractor.db.db.keys()) + ["other"]))
        # map skill -> category
        def cat_of(s):
            return extractor.db.category(s)
        # build arrays
        resume_cat_scores = {c:[] for c in cats}
        jd_cat_scores = {c:[] for c in cats}
        # For a JD skill, take its best similarity from matrix
        sim = result["similarity_matrix"]
        r_sk = result["resume_unique"]
        j_sk = result["jd_unique"]
        if sim.size:
            for j_idx, j in enumerate(j_sk):
                best_sim = float(sim[:, j_idx].max()) if sim.shape[0] > 0 else 0.0
                jd_cat_scores[cat_of(j)].append(best_sim)
            for r_idx, r in enumerate(r_sk):
                # approximate resume coverage by best similarity to any JD skill
                best_sim = float(sim[r_idx, :].max()) if sim.shape[1] > 0 else 0.0
                resume_cat_scores[cat_of(r)].append(best_sim)
        # average per category
        categories = sorted(cats)
        resume_vals = [np.mean(resume_cat_scores[c]) if len(resume_cat_scores[c])>0 else 0.0 for c in categories]
        jd_vals = [np.mean(jd_cat_scores[c]) if len(jd_cat_scores[c])>0 else 0.0 for c in categories]
        # radar requires closed loop
        fig = go.Figure()
        fig.add_trace(go.Scatterpolar(r=resume_vals + [resume_vals[0]], theta=categories + [categories[0]], fill='toself', name='Resume'))
        fig.add_trace(go.Scatterpolar(r=jd_vals + [jd_vals[0]], theta=categories + [categories[0]], fill='toself', name='Job Description'))
        fig.update_layout(polar=dict(radialaxis=dict(range=[0,1])), showlegend=True, title="Category-wise Coverage Radar (0-1)")
        return fig

# ---------------- Reports / Learning path ----------------

class ReportGenerator:
    def __init__(self):
        self.timestamp = datetime.now()

    def generate_text(self, analysis: Dict[str,Any]) -> str:
        lines = []
        lines.append("AI Skill Gap Analysis Report")
        lines.append(f"Generated: {self.timestamp.isoformat()}")
        lines.append("")
        lines.append(f"Overall match: {analysis.get('overall', analysis.get('overall_score',0.0))*100:.1f}%")
        lines.append(f"Matched: {len(analysis['matched'])} | Partial: {len(analysis['partial'])} | Missing: {len(analysis['missing'])}")
        lines.append("")
        lines.append("Top matched:")
        for m in analysis['matched'][:20]:
            lines.append(f" - {m['jd_skill']} -> {m['resume_skill']} ({m['similarity']*100:.1f}%)")
        lines.append("")
        lines.append("Top missing:")
        for m in sorted(analysis['missing'], key=lambda x: (1-x['similarity']))[:50]:
            lines.append(f" - {m['jd_skill']} ({m['similarity']*100:.1f}%)")
        return "\n".join(lines)

    def generate_csv(self, analysis: Dict[str,Any]) -> str:
        rows = []
        for x in analysis['matched']:
            rows.append({"JD Skill": x['jd_skill'], "Resume Match": x['resume_skill'], "Similarity": f"{x['similarity']*100:.2f}", "Status":"Matched"})
        for x in analysis['partial']:
            rows.append({"JD Skill": x['jd_skill'], "Resume Match": x['resume_skill'], "Similarity": f"{x['similarity']*100:.2f}", "Status":"Partial"})
        for x in analysis['missing']:
            rows.append({"JD Skill": x['jd_skill'], "Resume Match": x['resume_skill'], "Similarity": f"{x['similarity']*100:.2f}", "Status":"Missing"})
        df = pd.DataFrame(rows)
        return df.to_csv(index=False)

    def generate_json(self, analysis: Dict[str,Any]) -> str:
        return json.dumps(analysis, indent=2, default=str)

    def generate_pdf(self, analysis: Dict[str,Any], filename="skillgap_report.pdf") -> str:
        if not reportlab_available:
            raise RuntimeError("reportlab not installed")
        c = canvas.Canvas(filename, pagesize=letter)
        text = c.beginText(40, 750)
        text.setFont("Helvetica", 12)
        text.textLine("AI Skill Gap Analysis Report")
        text.textLine(f"Generated: {self.timestamp.isoformat()}")
        text.textLine("")
        text.textLine(f"Overall Match: {analysis.get('overall',0.0)*100:.1f}%")
        c.drawText(text)
        c.showPage()
        c.save()
        return filename

class LearningPathGenerator:
    def __init__(self):
        self.db = {
            "Python": {"time":"4-8 weeks", "resources":["Python for Everybody","Automate the Boring Stuff"]},
            "SQL": {"time":"2-4 weeks", "resources":["SQLBolt","Mode SQL Tutorial"]},
            "TensorFlow": {"time":"6-12 weeks", "resources":["Coursera DL","TensorFlow docs"]},
            "AWS": {"time":"6-12 weeks", "resources":["AWS Cloud Practitioner","Hands-on labs"]}
        }

    def generate(self, missing: List[Dict[str,Any]], current_skills: List[str]) -> List[Dict[str,Any]]:
        plan = []
        for item in sorted(missing, key=lambda x: x['similarity']):
            sk = item['jd_skill']
            info = self.db.get(sk, None)
            plan_item = {"skill": sk, "current_similarity": item['similarity'], "priority": "HIGH" if item['similarity'] < 0.4 else "MEDIUM", "time_estimate": info['time'] if info else "Varies", "resources": info['resources'] if info else [f"Search for {sk} courses online"]}
            plan.append(plan_item)
        return plan

# -------------------- Streamlit App --------------------

def main():
    st.sidebar.title("Configuration & Quick Actions")
    st.sidebar.markdown("Select model, thresholds, and controls")

    model_name = st.sidebar.selectbox("Sentence-BERT model", ["all-MiniLM-L6-v2"], index=0)
    strong_thr = st.sidebar.slider("Strong match threshold", 0.5, 0.95, 0.80, 0.05)
    partial_thr = st.sidebar.slider("Partial match threshold", 0.2, 0.7, 0.50, 0.05)
    conf_filter = st.sidebar.slider("Skill confidence filter", 0.0, 1.0, 0.4, 0.05)
    st.sidebar.markdown("---")
    st.sidebar.markdown("Uploads: PDF, DOCX, TXT. OCR requires pdf2image + poppler + pytesseract.")
    if st.sidebar.button("Clear session"):
        for k in list(st.session_state.keys()):
            del st.session_state[k]
        st.experimental_rerun()

    # Top: Dual upload areas
    st.header("1 — Upload Documents")
    left, right = st.columns(2)
    with left:
        st.subheader("Upload Resumes")
        uploaded_resumes = st.file_uploader("Upload Resume files (PDF, DOCX, TXT)", type=["pdf","docx","txt"], accept_multiple_files=True, key="upload_resumes")
    with right:
        st.subheader("Upload Job Descriptions")
        uploaded_jds = st.file_uploader("Upload JD files (PDF, DOCX, TXT)", type=["pdf","docx","txt"], accept_multiple_files=True, key="upload_jds")

    # Info about OCR availability
    # if pytesseract_available:
        #st.info(f"pytesseract available (configured: {getattr(pytesseract,'pytesseract', 'PATH')})")
    #else:
        #st.warning("pytesseract not available: OCR disabled for scanned PDFs.")
    #if not pdf2image_available:
        #st.warning("pdf2image not installed: OCR for PDFs requires pdf2image + poppler.")

    processor = DocumentProcessor()
    results = []
    errors = []

    # process both sets and show per-file progress
    all_files = []
    if uploaded_resumes:
        all_files.extend([(f, "resume") for f in uploaded_resumes])
    if uploaded_jds:
        all_files.extend([(f, "jd") for f in uploaded_jds])

    if all_files:
        overall_progress = st.progress(0)
        status_placeholder = st.empty()
        total = len(all_files)
        for i, (f, ftype) in enumerate(all_files):
            status_placeholder.text(f"Processing {f.name} ({i+1}/{total})")
            valid, err = processor.validate_file(f)
            st.progress(int((i)/total*100))
            if not valid:
                errors.append({"name": f.name, "error": err})
                continue
            # per-file mini progress placeholder
            file_prog = st.progress(0)
            info = processor.extract_text(f)
            file_prog.progress(100)
            file_prog.empty()
            info["doc_type"] = ftype
            results.append(info)
            overall_progress.progress(int((i+1)/total*100))
        status_placeholder.empty()
        overall_progress.empty()
            # ---------------- Processing Results Summary ----------------
    # Show a compact summary card after processing completes
    if all_files is not None:
        total_docs = len(all_files)
        success_docs = len(results)
        failed_docs = len(errors)
        success_rate = (success_docs / total_docs) * 100 if total_docs else 0.0

        st.markdown("### 📊 Processing Results")

        # Three metrics in a row
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("📄 Total Documents", total_docs)
        with col2:
            st.metric("✅ Successfully Processed", success_docs)
        with col3:
            st.metric("❌ Failed", failed_docs)

        # Prominent success rate line
        st.markdown(f"<h4 style='color:#004AAD;'>🎯 Success Rate: {success_rate:.1f}%</h4>", unsafe_allow_html=True)

        # Friendly informational message depending on outcomes
        if total_docs and success_docs == total_docs and failed_docs == 0:
            st.success(f"✅ Processing complete! {success_docs}/{total_docs} documents processed successfully.")
        elif failed_docs > 0:
            st.warning(f"⚠️ Processing finished with errors: {failed_docs} failed out of {total_docs}. See 'Errors / Failed files' below.")
        else:
            st.info(f"Processing finished: {success_docs}/{total_docs} processed.")


    # show errors
    if errors:
        st.subheader("Errors / Failed files")
        for e in errors:
            st.error(f"{e['name']}: {e['error']}")

    if results:
        # show upload status table
        st.subheader("Upload Status")
        rows = []
        for r in results:
            rows.append({"filename": r["name"], "format": r["format"], "pages": r.get("pages",""), "tokens": r["tokens"], "status": r["parse_status"]})
        st.dataframe(pd.DataFrame(rows), use_container_width=True)

        # Preview docs and allow toggles per file
        st.subheader("2 — Preview & Normalize")
        for idx, r in enumerate(results):
            with st.expander(f"{r['name']} — {r['format']} — {r['parse_status']}"):
                c1, c2 = st.columns([3,1])
                with c1:
                    show_raw = st.checkbox(f"Show raw ({r['name']})", value=False, key=f"raw_{idx}")
                    # highlighted view placeholder
                    if show_raw:
                        st.text_area("Raw Text", r['raw_text'][:200000], height=300)
                    else:
                        st.markdown(r['clean_text'][:200000])
                with c2:
                    # normalization chip
                    st.markdown(f"<div class='chip'>Tokens: {r['tokens']}</div>", unsafe_allow_html=True)
                    norm = r['normalization_summary']
                    st.json(norm)
                    if r.get('ocr_confidence') is not None:
                        st.metric("OCR Confidence", f"{r['ocr_confidence']*100:.1f}%")
                    if r.get('error'):
                        st.error(r['error'])

        # allow assignment of docs
        st.subheader("3 — Assign Document Roles")
        file_names = [r["name"] for r in results]
        resumes_sel = st.multiselect("Select which files are Resumes", file_names, default=[r["name"] for r in results if r["doc_type"]=="resume"])
        jds_sel = st.multiselect("Select which files are Job Descriptions", file_names, default=[r["name"] for r in results if r["doc_type"]=="jd"])

        # Combine texts
        resumes_texts = [r["clean_text"] for r in results if r["name"] in resumes_sel]
        jds_texts = [r["clean_text"] for r in results if r["name"] in jds_sel]

        if st.button("🚀 Extract Skills & Run Analysis"):
            if not resumes_texts or not jds_texts:
                st.error("Please select at least one resume and one job description.")
            else:
                with st.spinner("Extracting skills..."):
                    extractor = SkillExtractor(use_spacy=True, use_keyword=True, use_fuzzy=True)
                    resume_corpus = "\n\n".join(resumes_texts)
                    jd_corpus = "\n\n".join(jds_texts)
                    resume_sk_struct = extractor.extract(resume_corpus)
                    jd_sk_struct = extractor.extract(jd_corpus)

                    # Filter by confidence
                    resume_skills = [s["skill"] for s in resume_sk_struct if s["confidence"] >= conf_filter]
                    jd_skills = [s["skill"] for s in jd_sk_struct if s["confidence"] >= conf_filter]

                    # fallback to noun chunks if nothing found
                    if not resume_skills:
                        doc_r = extractor.nlp(resume_corpus)
                        resume_skills = list({chunk.text.strip() for chunk in doc_r.noun_chunks})[:50]
                    if not jd_skills:
                        doc_j = extractor.nlp(jd_corpus)
                        jd_skills = list({chunk.text.strip() for chunk in doc_j.noun_chunks})[:50]

                st.success(f"Extracted {len(resume_skills)} resume skills and {len(jd_skills)} JD skills (post-filter).")

                # embeddings + analysis
                st.info("Encoding skills and computing similarity (Sentence-BERT)...")
                embedder = SentenceBERTEmbedder(model_name)
                analyzer = SkillGapAnalyzer(embedder, strong=strong_thr, partial=partial_thr)
                analysis = analyzer.analyze(resume_skills, jd_skills)
                # attach extractor for radar & categories
                analysis["_extractor"] = extractor

                # session state
                st.session_state["analysis"] = analysis
                st.session_state["results"] = results
                st.session_state["resume_skills"] = resume_skills
                st.session_state["jd_skills"] = jd_skills

                st.success("Analysis complete!")

    # If analysis exists, show results
    if st.session_state.get("analysis"):
        analysis = st.session_state["analysis"]
        extractor = analysis.get("_extractor", SkillExtractor())
        st.header("4 — Results & Visualizations")
        tabs = st.tabs(["Overview","Documents","Skills Extracted","Gap Analysis","Similarity Matrix","Reports","Learning Path"])
        with tabs[0]:
            st.subheader("Overview")
            st.metric("Overall Match %", f"{analysis['overall']*100:.1f}%")
            st.write(f"Matched: {len(analysis['matched'])} | Partial: {len(analysis['partial'])} | Missing: {len(analysis['missing'])}")
            # pie
            fig_pie = Visualizer.pie(analysis)
            st.plotly_chart(fig_pie, use_container_width=True)
            # top missing
            st.subheader("Top missing (high priority)")
            top_missing = sorted(analysis["missing"], key=lambda x: (x["importance"]*(1-x["similarity"])))[:5]
            for m in top_missing:
                st.markdown(f"**{m['jd_skill']}** — similarity {m['similarity']*100:.1f}% — recommendation: take course / practice")
        with tabs[1]:
            st.subheader("Documents (Side-by-side)")
            docs = st.session_state["results"]
            combined_resume = "\n\n".join([r["clean_text"] for r in docs if r["name"] in resumes_sel])
            combined_jd = "\n\n".join([r["clean_text"] for r in docs if r["name"] in jds_sel])
            c1, c2 = st.columns(2)
            with c1:
                st.subheader("Resumes (combined)")
                st.text_area("Resume text", combined_resume[:200000], height=300)
            with c2:
                st.subheader("Job Descriptions (combined)")
                st.text_area("JD text", combined_jd[:200000], height=300)
        with tabs[2]:
            st.subheader("Skills Extracted")
            # build skill structures to show counts
            resume_struct = extractor.extract("\n\n".join(resumes_texts)) if resumes_texts else []
            jd_struct = extractor.extract("\n\n".join(jds_texts)) if jds_texts else []
            st.markdown("**Resume skills (filtered)**")
            for s in st.session_state["resume_skills"]:
                st.markdown(f"<span class='skill-btn skill-strong'>{s}</span>", unsafe_allow_html=True)
            st.markdown("**JD skills (filtered)**")
            for s in st.session_state["jd_skills"]:
                st.markdown(f"<span class='skill-btn skill-missing'>{s}</span>", unsafe_allow_html=True)
            # --- Top Resume and JD Skills in button form (no tag cloud) ---
            st.markdown("### 🧾 Top Resume Skills")
            if "resume_skills" in st.session_state and st.session_state["resume_skills"]:
              top_resume = st.session_state["resume_skills"][:20]  # top 20
            for s in top_resume:
                st.markdown(f"<span class='skill-btn skill-strong'>{s}</span>", unsafe_allow_html=True)
            else:
                st.info("No resume skills found.")
            st.markdown("### 📄 Top Job Description Skills")
            if "jd_skills" in st.session_state and st.session_state["jd_skills"]:
                top_jd = st.session_state["jd_skills"][:20]  # top 20
                for s in top_jd:
                    st.markdown(f"<span class='skill-btn skill-missing'>{s}</span>", unsafe_allow_html=True)
                    
            # confidence histogram per category
            if resume_struct:
                df_r = pd.DataFrame(resume_struct)
                st.subheader("Resume - Confidence by Category")
                st.plotly_chart(px.histogram(df_r, x="category", y="confidence", histfunc="avg", title="Avg Confidence per Category"), use_container_width=True)
        with tabs[3]:
            st.subheader("Gap Analysis")
            st.markdown("**Strong Matches**")
            for m in analysis["matched"]:
                st.markdown(f"<span class='skill-btn skill-strong'>{m['jd_skill']} ↔ {m['resume_skill']} ({m['similarity']*100:.1f}%)</span>", unsafe_allow_html=True)
            st.markdown("**Partial Matches**")
            for m in analysis["partial"]:
                st.markdown(f"<span class='skill-btn skill-partial'>{m['jd_skill']} ↔ {m['resume_skill']} ({m['similarity']*100:.1f}%)</span>", unsafe_allow_html=True)
            st.markdown("**Missing Skills** (mark covered if transferable)")
            # allow marking covered
            missing_skills = [m["jd_skill"] for m in analysis["missing"]]
            if missing_skills:
                covered = st.multiselect("Mark missing skills as covered by transferable experience (these will be excluded from missing list)", missing_skills, key="covered")
                # recompute effective missing/overall if covered selected
                effective_missing = [m for m in analysis["missing"] if m["jd_skill"] not in covered]
                effective_matched = analysis["matched"] + [m for m in analysis["missing"] if m["jd_skill"] in covered]
            else:
                effective_missing = []
                effective_matched = analysis["matched"]
            # show lists
            for m in analysis["missing"]:
                if m["jd_skill"] in covered:
                    st.markdown(f"<span class='skill-btn skill-partial'>{m['jd_skill']} — marked covered (closest: {m['resume_skill']})</span>", unsafe_allow_html=True)
                else:
                    st.markdown(f"<span class='skill-btn skill-missing'>{m['jd_skill']}</span> — closest: {m['resume_skill']} ({m['similarity']*100:.1f}%)", unsafe_allow_html=True)
            # bar chart for missing
            fig_bar = Visualizer.bar_missing(analysis)
            st.plotly_chart(fig_bar, use_container_width=True)
            # Radar chart (category wise)
            fig_radar = Visualizer.radar_category(analysis, analysis["resume_unique"], analysis["jd_unique"], extractor)
            st.plotly_chart(fig_radar, use_container_width=True)
        with tabs[4]:
            st.subheader("Similarity Matrix")
            sim = analysis["similarity_matrix"]
            if sim.size:
                fig_heat = Visualizer.heatmap(sim, analysis["resume_unique"], analysis["jd_unique"], top=40)
                st.plotly_chart(fig_heat, use_container_width=True)
                df_mat = pd.DataFrame(sim, index=analysis["resume_unique"], columns=analysis["jd_unique"])
                csv_mat = df_mat.to_csv()
                st.download_button("Download similarity matrix (CSV)", csv_mat, file_name=f"similarity_matrix_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv")
            else:
                st.info("No similarity matrix available")
        with tabs[5]:
            st.subheader("Reports & Export")
            rg = ReportGenerator()
            text_rep = rg.generate_text(analysis)
            csv_rep = rg.generate_csv(analysis)
            json_rep = rg.generate_json(analysis)
            st.download_button("Download TXT report", text_rep, file_name=f"skillgap_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
            st.download_button("Download CSV report", csv_rep, file_name=f"skillgap_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv")
            st.download_button("Download JSON data", json_rep, file_name=f"skillgap_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
            if reportlab_available:
                try:
                    pdf_path = rg.generate_pdf(analysis, filename=f"skillgap_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")
                    with open(pdf_path,"rb") as f:
                        st.download_button("Download PDF report", f, file_name=os.path.basename(pdf_path))
                except Exception as e:
                    st.error(f"PDF export failed: {e}")
            with tabs[6]:
              st.markdown("<h2 style='text-align:center; color:#004AAD;'>🎓 Personalized Learning Path</h2>", unsafe_allow_html=True)
              st.markdown("""
              This section provides a custom learning plan based on your missing skills from the analysis.
              Each skill is tagged with priority, difficulty level, and an estimated duration along with useful resources to improve quickly.
            """)

            lpg = LearningPathGenerator()
            plan = lpg.generate(analysis["missing"], analysis["resume_unique"])

    if plan:
        for i, p in enumerate(plan, 1):
            st.markdown(f"""
            <div style='background:linear-gradient(135deg,#e3f6ff,#ccecff);
                        padding:15px; border-radius:12px; margin-bottom:12px;
                        box-shadow:0px 3px 6px rgba(0,0,0,0.1);'>
                <h4 style='color:#004AAD;'>{i}. {p['skill']}</h4>
                <p><b>Priority:</b> {p['priority']} | <b>Level:</b> {p.get('level', 'Varies')} | 
                <b>Estimated Duration:</b> {p['time_estimate']}</p>
                <p><b>Recommended Resources:</b></p>
                <ul>
                    {''.join([f"<li>{r}</li>" for r in p['resources']])}
                </ul>
            </div>
            """, unsafe_allow_html=True)

             # ✅ Message appears ONLY once under this tab
            st.success("✅ Focus on **High Priority Beginner** skills first to strengthen your job readiness efficiently!")
        else:
                st.success("No missing skills — great match!")
    else:
        st.info("Upload resume(s) and job description(s) then click 'Extract Skills & Run Analysis' to start.")

if __name__ == "__main__":
    main() 